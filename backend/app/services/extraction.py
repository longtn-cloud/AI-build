import io

import docx
from pypdf import PdfReader


def extract_text(file_bytes: bytes, file_type: str) -> str:
    if file_type == "pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if file_type == "docx":
        document = docx.Document(io.BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n\n".join(parts)

    if file_type in ("txt", "md"):
        return file_bytes.decode("utf-8", errors="replace")

    raise ValueError(f"Unsupported file type: {file_type}")
