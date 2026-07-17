import io

import pytest
from docx import Document as DocxDocument

from app.services.extraction import extract_text


def _build_minimal_pdf(text: str) -> bytes:
    """Hand-build a minimal single-page PDF with an extractable text string.

    Byte-perfect xref offsets aren't required: pypdf falls back to
    startxref-scanning recovery when the xref table doesn't parse cleanly.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    stream_content = f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET".encode("latin-1")
    objects.append(
        b"<< /Length " + str(len(stream_content)).encode() + b" >>\nstream\n"
        + stream_content + b"\nendstream"
    )

    buffer = io.BytesIO()
    buffer.write(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(buffer.tell())
        buffer.write(f"{i} 0 obj\n".encode())
        buffer.write(obj)
        buffer.write(b"\nendobj\n")

    xref_offset = buffer.tell()
    count = len(objects) + 1
    buffer.write(f"xref\n0 {count}\n".encode())
    buffer.write(b"0000000000 65535 f \n")
    for offset in offsets:
        buffer.write(f"{offset:010d} 00000 n \n".encode())
    buffer.write(b"trailer\n")
    buffer.write(f"<< /Size {count} /Root 1 0 R >>\n".encode())
    buffer.write(b"startxref\n")
    buffer.write(f"{xref_offset}\n".encode())
    buffer.write(b"%%EOF")

    return buffer.getvalue()


def test_extracts_pdf_text():
    pdf_bytes = _build_minimal_pdf("Hello PDF World")

    result = extract_text(pdf_bytes, "pdf")

    assert "Hello PDF World" in result


def test_extracts_plain_text():
    result = extract_text(b"hello world", "txt")
    assert result == "hello world"


def test_extracts_markdown_as_plain_text():
    result = extract_text(b"# Title\n\nBody text", "md")
    assert result == "# Title\n\nBody text"


def test_extracts_docx_paragraphs():
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    buffer = io.BytesIO()
    doc.save(buffer)

    result = extract_text(buffer.getvalue(), "docx")

    assert "First paragraph" in result
    assert "Second paragraph" in result


def test_unsupported_type_raises():
    with pytest.raises(ValueError):
        extract_text(b"data", "exe")
